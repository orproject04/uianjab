import docx
import json

def extract_docx(filepath):
    doc = docx.Document(filepath)

    result = {
        "nama_jabatan": "",
        "unit_kerja": "",
        "ikhtisar_jabatan": "",
        "tugas_pokok": [],
        "jumlah_pegawai_dibutuhkan": "",
        "pembulatan": ""
    }

    # --- Scan metadata ---
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue

            key, val = None, None

            if len(cells) >= 3 and cells[1] == ":":
                key, val = cells[0], cells[2]
            elif len(cells) >= 2:
                key, val = cells[0], cells[1]
            elif len(cells) == 1 and ":" in cells[0]:
                parts = cells[0].split(":", 1)
                key, val = parts[0], parts[1]

            if not key:
                continue

            key_low = key.lower()
            if "nama jabatan" in key_low:
                result["nama_jabatan"] = val.strip()
            elif "unit kerja" in key_low:
                result["unit_kerja"] = val.strip()
            elif "ikhtisar jabatan" in key_low:
                result["ikhtisar_jabatan"] = val.strip()

    # --- Cari tabel tugas pokok ---
    for table in doc.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if any("uraian tugas" in h for h in headers):
            colmap = {}
            for idx, h in enumerate(headers):
                if "uraian tugas" in h: colmap["uraian_tugas"] = idx
                elif "satuan hasil" in h: colmap["satuan_hasil"] = idx
                elif "waktu penyelesaian" in h: colmap["waktu_penyelesaian"] = idx
                elif "waktu kerja" in h: colmap["waktu_kerja_efektif"] = idx
                elif "beban kerja" in h: colmap["beban_kerja"] = idx
                elif "pegawai" in h: colmap["pegawai_dibutuhkan"] = idx

            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue

                # --- Tangani baris khusus ---
                if any("Jumlah Pegawai" in c for c in cells):
                    result["jumlah_pegawai_dibutuhkan"] = cells[-1]
                    continue
                if any("Pembulatan" in c for c in cells):
                    result["pembulatan"] = cells[-1]
                    continue

                # --- Normal row ---
                uraian_text = cells[colmap["uraian_tugas"]]
                tahapan = []
                if "Tahapan:" in uraian_text:
                    parts = uraian_text.split("Tahapan:")
                    uraian_tugas = parts[0].strip(" :;\n")
                    tahapan = [t.strip() for t in parts[1].split("\n") if t.strip()]
                else:
                    uraian_tugas = uraian_text.strip()

                result["tugas_pokok"].append({
                    "uraian_tugas": uraian_tugas,
                    "tahapan": tahapan,
                    "satuan_hasil": cells[colmap.get("satuan_hasil", 0)],
                    "waktu_penyelesaian": cells[colmap.get("waktu_penyelesaian", 0)],
                    "waktu_kerja_efektif": cells[colmap.get("waktu_kerja_efektif", 0)],
                    "beban_kerja": cells[colmap.get("beban_kerja", 0)],
                    "pegawai_dibutuhkan": cells[colmap.get("pegawai_dibutuhkan", 0)]
                })

    return result

import sys

if __name__ == "__main__":
    file_path = globals().get("__file_path__", None)

    if not file_path:
        if len(sys.argv) >= 2:
            file_path = sys.argv[1]

    if file_path:
        try:
            data = extract_docx(file_path)
            print(json.dumps(data, ensure_ascii=False))
        except Exception as e:
            print(f"❌ Error: {str(e)}", file=sys.stderr)
            sys.exit(1)
    else:
        print("❌ Argumen tidak lengkap: butuh file .doc/.docx", file=sys.stderr)
        sys.exit(1)
