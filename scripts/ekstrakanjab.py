import os
import re
import json
import tempfile
import subprocess
from docx import Document
from docx.oxml.ns import qn

# -------------------- UTIL --------------------

def clean(text: str) -> str:
    # Samakan dengan versi Anda (hapus BEL \u0007, CR, TAB, VT, FF)
    return re.sub(r'[\u0007\r\t\x0b\x0c]', '', (text or '')).strip()

def is_list_paragraph(p) -> bool:
    """
    Deteksi bullet/numbering tanpa COM:
    Cek keberadaan w:numPr pada w:pPr.
    """
    try:
        pPr = p._p.pPr
        if pPr is None:
            return False
        numPr = pPr.find(qn('w:numPr'))
        return numPr is not None
    except Exception:
        return False

def para_text(p) -> str:
    return clean(p.text)

def cell_text(cell) -> str:
    # .text sudah gabungkan semua paragraph
    return clean(cell.text)

def table_header_cells(table):
    # Ambil row pertama sebagai header
    hdr = []
    if table.rows:
        for c in table.rows[0].cells:
            hdr.append(cell_text(c).lower())
    return hdr

def split_items(text: str):
    return [item.strip() for item in (text or '').split("|||") if item.strip()]

def iter_block_items(doc):
    """
    Iterasi block (paragraph/table) dalam urutan aslinya.
    Sumber: pola umum python-docx.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield ("p", Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield ("t", Table(child, doc))

# -------------------- READER --------------------

def read_docx(file_path):
    doc = Document(file_path)
    lines = [para_text(p) for p in doc.paragraphs if para_text(p)]
    return doc, lines

def convert_doc_to_docx_via_libreoffice(src_path: str) -> str:
    """
    Konversi .doc -> .docx via LibreOffice headless. Return path .docx di temp.
    """
    tmpdir = tempfile.mkdtemp(prefix="doc2docx_")
    cmd = [
        "soffice", "--headless", "--convert-to", "docx", src_path, "--outdir", tmpdir
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except Exception as e:
        raise RuntimeError(f"Gagal konversi .doc ke .docx dengan LibreOffice: {e}")
    base = os.path.splitext(os.path.basename(src_path))[0]
    out = os.path.join(tmpdir, base + ".docx")
    if not os.path.exists(out):
        raise RuntimeError("File hasil konversi .docx tidak ditemukan.")
    return out

def read_doc(file_path):
    # Gantikan COM: konversi ke .docx lalu baca dengan python-docx
    docx_path = convert_doc_to_docx_via_libreoffice(file_path)
    return read_docx(docx_path)

# -------------------- EXTRACTORS (preserve JSON shape) --------------------

def extract_line_value(label, lines):
    for line in lines:
        if label in line:
            parts = line.split(":")
            if len(parts) >= 2:
                return parts[1].strip()
    return "---"

def extract_unit_kerja(lines):
    unit_kerja = {
        "JPT Utama": "---", "JPT Madya": "---", "JPT Pratama": "---",
        "Administrator": "---", "Pengawas": "---", "Pelaksana": "---",
        "Jabatan Fungsional": "---"
    }
    start = False
    for line in lines:
        if "UNIT KERJA" in line.upper():
            start = True
            continue
        if start:
            if "IKHTISAR JABATAN" in line.upper() or "KUALIFIKASI JABATAN" in line.upper():
                break
            for key in unit_kerja:
                if key.upper() in line.upper():
                    value = line.split(":")[-1].strip()
                    unit_kerja[key] = value
    return unit_kerja

def extract_block(start_marker, end_marker, lines):
    start_idx = end_idx = None
    for i, line in enumerate(lines):
        if start_marker in line:
            start_idx = i + 1
        elif end_marker in line and start_idx is not None:
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        return "\n".join(lines[start_idx:end_idx]).strip()
    return "---"

def extract_kualifikasi(doc):
    """
    Baca tabel KUALIFIKASI JABATAN (3 kolom, 6+ baris) dengan dukungan alias label:
    - 'Pendidikan Formal'         -> pendidikan_formal
    - 'Diklat Penjenjangan' OR 'Manajerial' (juga 'Penjenjangan', 'Diklat Manajerial', ...) -> diklat_penjenjangan
    - 'Diklat Teknis' OR 'Teknis' -> diklat_teknis
    - 'Diklat Fungsional' OR 'Fungsional' -> diklat_fungsional
    - 'Pengalaman Kerja'          -> pengalaman_kerja
    """
    result = {
        "pendidikan_formal": [],
        "pendidikan_dan_pelatihan": {
            "diklat_penjenjangan": [],
            "diklat_teknis": [],
            "diklat_fungsional": []
        },
        "pengalaman_kerja": []
    }

    def tidy(s: str) -> str:
        s = (s or "").replace("\u00A0", " ")
        s = re.sub(r"[\u0000-\u001F]", "", s)
        s = s.strip().strip(":;")
        # buang bullet/nomor di depan
        s = re.sub(r"^\s*(?:[\u2022\-\–\—\*•]+|\(?\d+[\.\)])\s*", "", s)
        return s.strip()

    def cell_to_list(cell):
        items = []
        for p in cell.paragraphs:
            t = tidy(p.text)
            if not t or t in {"-", "–", "—", ":"}:
                continue
            items.append(t)
        if items:
            return items
        txt = tidy(cell.text)
        return [x for x in (ln.strip() for ln in txt.splitlines()) if x]

    # ===== Alias sets (lowercase) =====
    PENJENJANGAN_ALIASES = (
        "diklat penjenjangan", "penjenjangan", "manajerial", "diklat manajerial",
        "pelatihan manajerial"
    )
    TEKNIS_ALIASES = (
        "diklat teknis", "teknis", "pelatihan teknis"
    )
    FUNGSIONAL_ALIASES = (
        "diklat fungsional", "fungsional", "pelatihan fungsional"
    )

    # Untuk scoring kandidat tabel, gabungkan semua label yang kita kenal
    LABELS_FOR_SCORE = (
        "pendidikan formal",
        *PENJENJANGAN_ALIASES,
        *TEKNIS_ALIASES,
        *FUNGSIONAL_ALIASES,
        "pengalaman kerja",
    )

    def table_score(tbl):
        if len(tbl.rows) < 3:  # longgar: kadang >6, kadang <6
            return -1
        # butuh >=3 kolom di beberapa baris awal
        probe = min(6, len(tbl.rows))
        if any(len(r.cells) < 3 for r in tbl.rows[:probe]):
            return -1
        labels = [tidy(r.cells[0].text).lower() for r in tbl.rows[:probe]]
        return sum(any(lab in x for lab in LABELS_FOR_SCORE) for x in labels)

    candidate = None
    best = -1
    for t in doc.tables:
        sc = table_score(t)
        if sc > best:
            best = sc
            candidate = t

    if not candidate or best < 2:
        return result

    # ===== Helper matcher =====
    def match_any(label_low: str, aliases: tuple[str, ...]) -> bool:
        return any(a in label_low for a in aliases)

    # mapping per baris berdasarkan isi kolom-1
    for row in candidate.rows:
        if len(row.cells) < 3:
            continue
        c1_low = tidy(row.cells[0].text).lower()
        # Skip baris header "Pendidikan dan Pelatihan"
        if "pendidikan dan pelatihan" in c1_low:
            continue

        val_items = cell_to_list(row.cells[2])
        if not val_items:
            continue

        if "pendidikan formal" in c1_low:
            result["pendidikan_formal"] = val_items

        elif match_any(c1_low, PENJENJANGAN_ALIASES):
            result["pendidikan_dan_pelatihan"]["diklat_penjenjangan"] = val_items

        elif match_any(c1_low, TEKNIS_ALIASES):
            result["pendidikan_dan_pelatihan"]["diklat_teknis"] = val_items

        elif match_any(c1_low, FUNGSIONAL_ALIASES):
            result["pendidikan_dan_pelatihan"]["diklat_fungsional"] = val_items

        elif "pengalaman kerja" in c1_low:
            result["pengalaman_kerja"] = val_items

    return result


def extract_bullet_marked_text_cell(cell) -> str:
    """
    Replikasi versi COM:
    - Jika paragraf ber-numbering/bullet → prefix '|||'
    - Jika bukan → langsung gabung.
    """
    chunks = []
    for p in cell.paragraphs:
        t = para_text(p)
        if not t:
            continue
        if is_list_paragraph(p):
            chunks.append(f"|||{t}")
        else:
            chunks.append(t)
    return " ".join(chunks)

def extract_deskripsi_dan_tahapan_cell(cell):
    deskripsi = ""
    tahapan = []
    current_is_tahapan = False
    for p in cell.paragraphs:
        text = para_text(p)
        if not text:
            continue
        if "tahapan" in text.lower().strip(":").strip():
            current_is_tahapan = True
            continue
        if not current_is_tahapan:
            deskripsi += (" " if deskripsi else "") + text
        else:
            tahapan.append(text)
    return deskripsi.strip(), tahapan

# ---------- Bagian tabel: gunakan table_header_cells() untuk map kolom ----------

def extract_tugas_pokok(doc):
    """
    Ekstrak tabel TUGAS POKOK menjadi list item:
    - deskripsi (teks uraian)
    - detail_uraian_tugas (tahapan + detail)
    - hasil_kerja (bullet/daftar)
    - jumlah_hasil, waktu_penyelesaian_(jam), waktu_efektif, kebutuhan_pegawai (string kosong default)
    Catatan:
    - Mengabaikan baris total/rekap: "JUMLAH", "JUMLAH PEGAWAI", "PEMBULATAN".
    - Tidak mengabaikan baris biasa hanya karena mengandung kata 'jumlah' di tengah kalimat.
    """
    import re
    from docx.oxml.ns import qn

    # ---------- util dasar ----------
    def tidy(s: str) -> str:
        # buang kontrol & whitespace berlebih
        return re.sub(r'[\u0007\r\t\x0b\x0c]', '', (s or '')).strip()

    # pola penomoran
    NUMERIC_TOP_RE = re.compile(r'^\s*\d{1,3}[\.\)]\s+')
    LETTER_SUB_RE  = re.compile(r'^\s*[a-zA-Z][\.\)]\s+')

    # ---------- akses numbering Word ----------
    def get_numpr(p):
        try:
            pPr = p._p.pPr
            if pPr is None:
                return None, None
            numPr = pPr.find(qn('w:numPr'))
            if numPr is None:
                return None, None
            numId_elm = numPr.find(qn('w:numId'))
            ilvl_elm  = numPr.find(qn('w:ilvl'))
            numId = int(numId_elm.val) if (numId_elm is not None and numId_elm.val is not None) else None
            ilvl  = int(ilvl_elm.val)  if (ilvl_elm  is not None and ilvl_elm.val  is not None) else 0
            return numId, ilvl
        except Exception:
            return None, None

    def get_numfmt(doc, numId, ilvl):
        try:
            numbering = doc.part.numbering_part.element
            for num in numbering.findall(qn('w:num')):
                if str(num.get(qn('w:numId'))) == str(numId):
                    absId_elm = num.find(qn('w:abstractNumId'))
                    if absId_elm is None:
                        break
                    absId = absId_elm.get(qn('w:val'))
                    for absnum in numbering.findall(qn('w:abstractNum')):
                        if str(absnum.get(qn('w:abstractNumId'))) == str(absId):
                            # cari level tepat
                            for lvl in absnum.findall(qn('w:lvl')):
                                if str(lvl.get(qn('w:ilvl'))) == str(ilvl):
                                    nf = lvl.find(qn('w:numFmt'))
                                    if nf is not None and nf.get(qn('w:val')):
                                        return nf.get(qn('w:val')).lower()
                            # fallback lvl=0
                            for lvl in absnum.findall(qn('w:lvl')):
                                if str(lvl.get(qn('w:ilvl'))) == "0":
                                    nf = lvl.find(qn('w:numFmt'))
                                    if nf is not None and nf.get(qn('w:val')):
                                        return nf.get(qn('w:val')).lower()
                    break
        except Exception:
            pass
        return None

    def indent_twips(p):
        try:
            pf = p.paragraph_format
            left  = pf.left_indent.twips if pf.left_indent else 0
            first = pf.first_line_indent.twips if pf.first_line_indent else 0
            hang  = -first if (first is not None and first < 0) else 0
            return int(left + hang)
        except Exception:
            return 0

    SUB_DELTA = 120  # twips

    # ---------- klasifikasi baris uraian ----------
    def classify_line(doc, p, base_indent):
        """
        return "top" (tahapan) atau "sub" (detail_tahapan).
        PRIORITAS:
        1) Regex teks (angka = top, huruf = sub)
        2) numFmt (decimal-family = top; letter-family = sub)
        3) ilvl
        4) delta indent
        """
        txt = tidy(p.text)
        if not txt:
            return None, ""

        # (1) Regex teks
        if NUMERIC_TOP_RE.match(txt):
            return "top", txt
        if LETTER_SUB_RE.match(txt):
            return "sub", txt

        # (2) numFmt Word
        numId, ilvl = get_numpr(p)
        numfmt = get_numfmt(doc, numId, ilvl) if numId is not None else None
        if numfmt:
            # anggap semua decimal/roman angka sbg top
            if numfmt in {"decimal", "decimalzero", "arabic", "lowerroman", "upperroman"}:
                return "top", txt
            if numfmt in {"lowerletter", "loweralpha", "alphalower", "upperletter", "upperalpha", "alphaupper"}:
                return "sub", txt

        # (3) ilvl
        if ilvl is not None:
            if ilvl == 0:
                ind = indent_twips(p)
                if base_indent is not None and (ind - base_indent) >= SUB_DELTA:
                    return "sub", txt
                return "top", txt
            else:
                return "sub", txt

        # (4) Fallback indent
        ind = indent_twips(p)
        if base_indent is not None and (ind - base_indent) >= SUB_DELTA:
            return "sub", txt
        return "top", txt

    # ---------- parsing uraian -> deskripsi & nested ----------
    def parse_uraian_cell_to_struct(cell):
        deskripsi_parts = []
        detail_uraian   = []
        in_tahapan      = False
        current         = None
        base_indent     = None
        seen_any_line   = False

        for p in cell.paragraphs:
            raw = tidy(p.text)
            if not raw:
                continue

            low = raw.lower().strip().strip(':')
            if not in_tahapan and 'tahapan' in low:
                in_tahapan = True
                continue

            if not in_tahapan:
                deskripsi_parts.append(raw)
                continue

            kind, text = classify_line(doc, p, base_indent)
            if not kind:
                continue

            # Promosikan baris pertama bila terklasifikasi "sub"
            if not seen_any_line and kind == "sub":
                kind = "top"
            seen_any_line = True

            if kind == "top":
                if current:
                    detail_uraian.append(current)
                current = {"tahapan": text, "detail_tahapan": []}
                if base_indent is None:
                    base_indent = indent_twips(p)
            else:
                if current is None:
                    current = {"tahapan": "", "detail_tahapan": []}
                current["detail_tahapan"].append(text)

        if current:
            detail_uraian.append(current)

        return " ".join(deskripsi_parts).strip(), detail_uraian

    # # ---------- hasil kerja (pakai helper Anda yang sudah ada) ----------
    # def extract_bulleted_items(cell):
    #     raw = extract_bullet_marked_text_cell(cell)  # helper existing
    #     return split_items(raw)  # helper existing

    # ================== POLA TEKS ==================
    ALPHA_PREFIX   = re.compile(r"^\s*[A-Za-z][\.\)]\s")           # a.  a)
    DECIMAL_PREFIX = re.compile(r"^\s*\d+[\.\)]\s")                # 1.  1)
    ROMAN_PREFIX   = re.compile(r"^\s*[IVXLCDM]+[\.\)]\s", re.I)   # I.  iv)
    BULLET_LIKE    = re.compile(r"^\s*[•▪◦·\-–—■□]\s*")            # simbol bullet umum

    # ================== OPSIONAL: para_text yang membaca <w:br/> ==================
    # Aktifkan jika Anda ingin line-break Word (Shift+Enter) terbaca sebagai '\n'
    # lalu pecah menjadi "paragraf logis" terpisah sebelum diparse.
    # def para_text_with_br(p):
    #     texts = []
    #     for r in p.runs:
    #         # deteksi <w:br/>
    #         brs = r._element.findall(".//w:br", namespaces=r._element.nsmap)
    #         if brs:
    #             texts.append(r.text or "")
    #             texts.append("\n")
    #         else:
    #             texts.append(r.text or "")
    #     return "".join(texts).strip()

    # ================== PUBLIC API ==================
    def extract_bulleted_items(cell):
        """
        Menghasilkan struktur hasil_kerja -> List[{text, children[]}].
        - Mode A (abc): list huruf level-0 => parent; list beda gaya/level => child
        - Mode B (fallback: colon-block + title-streak + list-aware)
        """
        paras = []
        saw_alpha_parent = False

        for p in cell.paragraphs:
            # Jika pakai versi opsional:
            # t = (para_text_with_br(p) or "").strip()
            t = (para_text(p) or "").strip()
            if not t:
                continue

            # Jika Anda mengaktifkan para_text_with_br di atas, pecah '\n' jadi baris logis
            if "\n" in t:
                for part in [x.strip() for x in t.split("\n") if x.strip()]:
                    paras.append({"text": part, "is_list": False, "ilvl": 0, "numfmt0": None, "absId": None})
                continue

            meta = {"text": t, "is_list": False, "ilvl": 0, "numfmt0": None, "absId": None}

            if is_list_paragraph(p):
                meta["is_list"] = True
                try:
                    # level
                    pPr = p._p.pPr
                    numPr = pPr.numPr if pPr is not None else None
                    if numPr is not None and numPr.ilvl is not None and numPr.ilvl.val is not None:
                        meta["ilvl"] = int(numPr.ilvl.val)

                    # signature level-0
                    numId = int(numPr.numId.val) if (numPr is not None and numPr.numId is not None and numPr.numId.val is not None) else None
                    if numId is not None:
                        numbering = p.part.numbering_part.element
                        num = numbering.xpath(f".//w:num[@w:numId='{numId}']", namespaces=numbering.nsmap)
                        if num:
                            absId_el = num[0].xpath("./w:abstractNumId", namespaces=numbering.nsmap)
                            if absId_el:
                                meta["absId"] = absId_el[0].get(qn("w:val"))
                                absn = numbering.xpath(f".//w:abstractNum[@w:abstractNumId='{meta['absId']}']", namespaces=numbering.nsmap)
                                if absn:
                                    lvl0 = absn[0].xpath("./w:lvl[@w:ilvl='0']", namespaces=numbering.nsmap)
                                    if lvl0:
                                        numFmt_el = lvl0[0].xpath("./w:numFmt", namespaces=numbering.nsmap)
                                        if numFmt_el:
                                            meta["numfmt0"] = numFmt_el[0].get(qn("w:val"))
                except Exception:
                    pass

                if meta["ilvl"] == 0 and meta["numfmt0"] in ("lowerLetter", "upperLetter"):
                    saw_alpha_parent = True

            paras.append(meta)

        if saw_alpha_parent:
            return _build_mode_abc(paras)
        else:
            return _build_mode_colon_block(paras)

    # ================== MODE A: abc signature ==================
    def _build_mode_abc(paras):
        root = []
        current_parent = None
        base_sig = None  # (absId, numfmt0)

        for r in paras:
            t = r["text"]
            if r["is_list"]:
                sig = (r["absId"], r["numfmt0"])
                # inisialisasi base signature di list huruf level-0
                if base_sig is None and r["ilvl"] == 0 and r["numfmt0"] in ("lowerLetter", "upperLetter"):
                    base_sig = sig
                    current_parent = {"text": t, "children": []}
                    root.append(current_parent)
                    continue

                if base_sig is not None and r["ilvl"] == 0 and _same_sig(sig, base_sig):
                    # signature sama dengan base -> parent baru
                    current_parent = {"text": t, "children": []}
                    root.append(current_parent)
                    continue

                # signature beda atau ilvl>=1 -> child
                if current_parent is None:
                    current_parent = {"text": "(parent otomatis)", "children": []}
                    root.append(current_parent)
                current_parent["children"].append({"text": t, "children": []})
                continue

            # plain text
            if current_parent is None:
                current_parent = {"text": t, "children": []}
                root.append(current_parent)
            else:
                if _looks_childish(t):
                    current_parent["children"].append({"text": t, "children": []})
                else:
                    current_parent["text"] = (current_parent["text"].rstrip() + " " + t.lstrip()).strip()

        return root

    def _same_sig(a, b):
        if not a or not b:
            return False
        absA, fmtA = a
        absB, fmtB = b
        if absA and absB:
            return absA == absB
        if fmtA and fmtB:
            return fmtA == fmtB
        return False

    # ================== MODE B: colon-block + title-streak + list-aware ==================
    def _build_mode_colon_block(paras):
        """
        Parent:
        - baris berakhir ':'
        - inline 'Parent: Child...' → parent + anak (split cerdas pada tail)
        - (title-streak) bila parent pertama tanpa ':' dan tampak 'judul' (Kapital & >=3 kata),
            maka paragraf berikutnya yang juga 'judul' → parent sejajar
        - NEW: list level-0 dengan signature sama (absId, numfmt0) → parent sejajar (jika TIDAK dalam colon-parent)
        Children:
        - LIST Word (dalam colon-parent) → child
        - list level>=1 → child
        - list level-0 dgn signature berbeda dari base → child
        - numeric/bullet (teks) → child
        - short connector (≤2 kata, atau 'dan/atau/serta/hingga/maupun') → ditempel
        Pemutus blok:
        - jika child terakhir berakhir '.', paragraf biasa berikutnya (bukan angka/bullet/':') → parent baru
        """
        root = []
        current_parent = None
        last_child_ended_with_dot = False
        title_streak = False

        # NEW: track “apakah parent aktif berasal dari colon” dan “signature list L0”
        colon_parent_active = False
        base_list_sig = None  # (absId, numfmt0) untuk list level-0 di luar colon-parent

        def _is_title_like(t: str) -> bool:
            if not t:
                return False
            w = t.strip().split()
            return t[0].isupper() and len(w) >= 3

        def _is_short_connector(t: str) -> bool:
            w = t.strip().split()
            if len(w) < 3 and not t.rstrip().endswith((".", ":", ";")):
                return True  # termasuk 'dan/atau/serta/hingga/maupun'
            return False

        def start_parent(text, from_colon=False, sig=None):
            nonlocal current_parent, last_child_ended_with_dot, title_streak, colon_parent_active, base_list_sig
            node = {"text": text.strip(), "children": []}
            # catat metadata ringan (tidak ikut output)
            # node["_from_colon"] = bool(from_colon)
            # node["_sig"] = sig
            root.append(node)
            current_parent = node
            last_child_ended_with_dot = False
            colon_parent_active = bool(from_colon)
            # title-streak hanya aktif jika bukan dari colon
            title_streak = False if from_colon else _is_title_like(text)
            # jika ini parent list level-0 pertama di luar colon, set base signature
            if (sig is not None) and (not from_colon) and (base_list_sig is None):
                base_list_sig = sig
            return node

        def add_child(text):
            nonlocal current_parent, last_child_ended_with_dot
            if current_parent is None:
                start_parent("(parent otomatis)")
            current_parent["children"].append({"text": text.strip(), "children": []})
            last_child_ended_with_dot = text.rstrip().endswith(".")

        def append_to_last(text):
            nonlocal current_parent, last_child_ended_with_dot
            if current_parent["children"]:
                last = current_parent["children"][-1]
                last["text"] = (last["text"].rstrip() + " " + text.lstrip()).strip()
                last_child_ended_with_dot = last["text"].rstrip().endswith(".")
            else:
                current_parent["text"] = (current_parent["text"].rstrip() + " " + text.lstrip()).strip()
                last_child_ended_with_dot = False

        def _same_sig(a, b):
            if not a or not b:
                return False
            absA, fmtA = a
            absB, fmtB = b
            if absA and absB:
                return absA == absB
            if fmtA and fmtB:
                return fmtA == fmtB
            return False

        for r in paras:
            t = r["text"]

            # 0) Inline colon: "Parent: Child ... "
            if ":" in t and not t.strip().endswith(":"):
                head, tail = t.split(":", 1)
                head, tail = head.strip(), tail.strip()
                if head and tail:
                    start_parent(head + ":", from_colon=True, sig=None)
                    parts = _smart_split(tail)
                    if len(parts) <= 1:
                        parts = [p.strip() for p in re.split(r"(?<=\S)\s+(?=[A-Z])", tail) if p.strip()]
                    for seg in parts:
                        add_child(seg)
                    continue  # selesai baris ini

            # 1) Parent by colon di akhir
            if t.endswith(":"):
                start_parent(t, from_colon=True, sig=None)
                continue

            # 2) Pemutus blok titik → parent baru
            if last_child_ended_with_dot and not (
                DECIMAL_PREFIX.match(t) or ALPHA_PREFIX.match(t) or ROMAN_PREFIX.match(t) or BULLET_LIKE.match(t) or t.endswith(":")
            ):
                start_parent(t, from_colon=False, sig=None)
                continue

            # 3) Belum ada parent? jadikan parent awal (non-colon)
            if current_parent is None:
                # Jika ini list level-0, set sebagai parent dan simpan signature
                if r.get("is_list") and r.get("ilvl", 0) == 0:
                    sig = (r.get("absId"), r.get("numfmt0"))
                    start_parent(t, from_colon=False, sig=sig)
                else:
                    start_parent(t, from_colon=False, sig=None)
                continue

            # 3.5) LIST-AWARE DITINGKATKAN:
            if r.get("is_list"):
                ilvl = r.get("ilvl", 0)
                sig = (r.get("absId"), r.get("numfmt0"))

                if colon_parent_active:
                    # Sedang dalam colon-parent → semua list = child
                    add_child(t)
                    continue

                if ilvl == 0:
                    # NEW: kalau teksnya "judul", jadikan parent sejajar meski signature kosong
                    if _is_title_like(t):
                        start_parent(t, from_colon=False, sig=sig)
                        # set base_list_sig sekali kalau belum ada
                        if base_list_sig is None:
                            base_list_sig = sig
                        continue

                    if base_list_sig is None:
                        base_list_sig = sig
                        start_parent(t, from_colon=False, sig=sig)
                        continue

                    if _same_sig(sig, base_list_sig):
                        start_parent(t, from_colon=False, sig=sig)
                        continue
                    else:
                        add_child(t)
                        continue
                else:
                    # level dalam → child
                    add_child(t)
                    continue

            # 4) Title-streak: parent awal tanpa ':' dan ini 'judul' → parent sejajar
            if title_streak and _is_title_like(t) and not (
                DECIMAL_PREFIX.match(t) or ALPHA_PREFIX.match(t) or ROMAN_PREFIX.match(t) or BULLET_LIKE.match(t)
            ):
                start_parent(t, from_colon=False, sig=None)
                continue

            # 5) Numeric/bullet (teks) → child
            if DECIMAL_PREFIX.match(t) or ALPHA_PREFIX.match(t) or ROMAN_PREFIX.match(t) or BULLET_LIKE.match(t):
                add_child(t)
                continue

            # 6) Short connector / sambungan → ditempel
            if _is_short_connector(t):
                append_to_last(t)
                continue

            # 7) Default: tempel ke item terakhir (child jika ada, else ke parent)
            append_to_last(t)

        return root

    # ================== HEURISTIK TAMBAHAN ==================
    def _looks_childish(t: str) -> bool:
        return bool(
            BULLET_LIKE.match(t) or
            DECIMAL_PREFIX.match(t) or
            ALPHA_PREFIX.match(t) or
            ROMAN_PREFIX.match(t)
        )

    def _smart_split(text: str):
        """
        Pecah string pada:
        - koma di LUAR kurung
        - lalu ' dan ' di LUAR kurung
        - jika masih 1 segmen panjang, pecah pada huruf kapital di awal segmen berikutnya
        Hindari pecah di dalam '(...)'.
        """
        # 1) split by comma outside parentheses
        parts = []
        buf = []
        depth = 0
        for ch in text:
            if ch == '(':
                depth += 1; buf.append(ch)
            elif ch == ')':
                depth = max(0, depth - 1); buf.append(ch)
            elif ch == ',' and depth == 0:
                parts.append(''.join(buf).strip()); buf = []
            else:
                buf.append(ch)
        if buf:
            parts.append(''.join(buf).strip())

        # 2) split by ' dan ' outside parentheses
        if len(parts) == 1:
            s = parts[0]
            res, buf, depth, i = [], [], 0, 0
            while i < len(s):
                ch = s[i]
                if ch == '(':
                    depth += 1; buf.append(ch); i += 1
                elif ch == ')':
                    depth = max(0, depth - 1); buf.append(ch); i += 1
                elif depth == 0 and s[i:i+5].lower() == " dan ":
                    res.append(''.join(buf).strip()); buf = []; i += 5
                else:
                    buf.append(ch); i += 1
            tail = ''.join(buf).strip()
            if tail:
                res.append(tail)
            parts = [p for p in res if p]

        # 3) fallback split on Capital starts (untuk kasus tanpa tanda baca)
        if len(parts) == 1:
            s = parts[0]
            capital_split = [p.strip() for p in re.split(r"(?<=\S)\s+(?=[A-Z])", s) if p.strip()]
            if len(capital_split) > 1:
                parts = capital_split

        return [p for p in parts if p]



    # ---------- deteksi baris total/rekap yang harus di-skip ----------
    TOTAL_KEYS = {
        "jumlah", "jumlah pegawai", "jumlah pegawai yang dibutuhkan",
        "pembulatan"
    }

    def is_placeholder(s: str) -> bool:
        """Cek isian kosong/placeholder seperti titik2, dash, null, dsb."""
        low = s.lower().strip()
        return (
            low in {"", "-", "–", "—", "null"} or
            low.strip(".") == "" or
            all(ch == '.' for ch in low if ch != ' ')
        )

    def row_is_total_or_footer(cells):
        """
        True bila baris adalah 'JUMLAH', 'JUMLAH PEGAWAI', 'PEMBULATAN', dsb.
        Pengecekan dari seluruh sel, cukup 1 sel mengandung label kunci.
        """
        texts = [tidy(c.text) for c in cells]
        joined = " ".join(texts).lower()
        # cepat: kalau ada kata kunci utama, langsung anggap footer
        for k in TOTAL_KEYS:
            # cocokkan yang berdiri sendiri / di awal baris
            if re.search(rf'\b{k}\b', joined):
                # kecuali jika jelas bagian dari kalimat panjang (jarang terjadi untuk baris footer)
                # di case real, baris footer biasanya 1-2 kata saja.
                # Tambahan: jika hampir semua kolom lain placeholder, makin yakin footer
                non_ph = sum(0 if is_placeholder(t) else 1 for t in texts)
                if non_ph <= 2:
                    return True
        # Baris yang seluruh kolom placeholder & tanpa uraian berarti baris kosong → skip juga
        if all(is_placeholder(t) for t in texts):
            return True
        return False

    # ---------- proses utama ----------
    tugas_list = []

    for table in doc.tables:
        headers = table_header_cells(table)  # helper existing -> list lowercase header names
        if not headers:
            continue
        if ("uraian tugas" in headers) and ("hasil kerja" in headers):
            uraian_idx = headers.index("uraian tugas")
            hasil_idx  = headers.index("hasil kerja")

            # iterasi baris isi
            for r in table.rows[1:]:
                cells = r.cells
                if len(cells) <= max(uraian_idx, hasil_idx):
                    continue

                # Skip baris total/rekap/footer
                if row_is_total_or_footer(cells):
                    continue

                deskripsi, detail_uraian = parse_uraian_cell_to_struct(cells[uraian_idx])
                hasil_items = extract_bulleted_items(cells[hasil_idx])

                # Normalisasi & filter baris kosong
                norm_desc = tidy(deskripsi).strip(". -").strip()
                low_desc  = norm_desc.lower()

                # --- Skip hanya jika "JUMLAH" dsb sebagai label standalone,
                #     BUKAN jika kata 'jumlah' muncul sebagai bagian kalimat biasa.
                if low_desc in TOTAL_KEYS or re.fullmatch(r'jumlah\.?', low_desc):
                    continue

                # buang baris yang benar-benar kosong
                if low_desc in {"", "...........", ".........."}:
                    continue
                if not hasil_items:
                    raw_h = tidy(extract_bullet_marked_text_cell(cells[hasil_idx])).strip(". -").strip().lower()
                    if raw_h in {"", "...........", "..........", "-", "–", "—"}:
                        # seluruh kolom kanan kosong → kemungkinan bukan baris tugas
                        continue

                tugas_list.append({
                    "no": str(len(tugas_list) + 1),
                    "uraian_tugas": {
                        "deskripsi": norm_desc,
                        "detail_uraian_tugas": detail_uraian,
                        "hasil_kerja": hasil_items,
                        "jumlah_hasil": "",
                        "waktu_penyelesaian_(jam)": "",
                        "waktu_efektif": "",
                        "kebutuhan_pegawai": ""
                    }
                })

            # ---- tambahan safety: bila masih ada footer nyasar di ekor, drop trailing ----
            while tugas_list:
                tail = tugas_list[-1]["uraian_tugas"]["deskripsi"].strip().lower()
                if (tail in TOTAL_KEYS) or re.fullmatch(r'jumlah\.?', tail):
                    tugas_list.pop()
                else:
                    break

    return tugas_list

def extract_hasil_kerja(doc):
    hasil_list = []
    for table in doc.tables:
        headers = table_header_cells(table)
        if not headers:
            continue
        if ("hasil kerja" in headers) and ("satuan hasil" in headers):
            hasil_idx = headers.index("hasil kerja")
            satuan_idx = headers.index("satuan hasil")
            for r in table.rows[1:]:
                cells = r.cells
                if len(cells) <= max(hasil_idx, satuan_idx):
                    continue
                hasil_raw = extract_bullet_marked_text_cell(cells[hasil_idx])
                satuan_raw = extract_bullet_marked_text_cell(cells[satuan_idx])
                hasil_items = split_items(hasil_raw) or []
                satuan_items = split_items(satuan_raw) or []
                hasil_list.append({
                    "no": str(len(hasil_list) + 1),
                    "hasil_kerja": hasil_items,
                    "satuan_hasil": satuan_items
                })
    return hasil_list

def extract_bahan_kerja(doc):
    bahan_list = []
    for table in doc.tables:
        headers = table_header_cells(table)
        if not headers:
            continue
        if ("bahan kerja" in headers) and ("penggunaan dalam tugas" in headers):
            b_idx = headers.index("bahan kerja")
            p_idx = headers.index("penggunaan dalam tugas")
            for r in table.rows[1:]:
                cells = r.cells
                if len(cells) <= max(b_idx, p_idx):
                    continue
                bahan_items = split_items(extract_bullet_marked_text_cell(cells[b_idx])) or []
                penggunaan_items = split_items(extract_bullet_marked_text_cell(cells[p_idx])) or []
                bahan_list.append({
                    "no": str(len(bahan_list) + 1),
                    "bahan_kerja": bahan_items,
                    "penggunaan_dalam_tugas": penggunaan_items
                })
    return bahan_list

def extract_perangkat_kerja(doc):
    perangkat_list = []
    try:
        for table in doc.tables:
            headers = table_header_cells(table)
            if not headers:
                continue
            if ("perangkat kerja" in headers) and any("penggunaan" in h for h in headers):
                perangkat_idx = headers.index("perangkat kerja")
                penggunaan_idx = next((i for i, h in enumerate(headers) if "penggunaan" in h), None)
                for r in table.rows[1:]:
                    cells = r.cells
                    perangkat_raw = extract_bullet_marked_text_cell(cells[perangkat_idx])
                    penggunaan_raw = extract_bullet_marked_text_cell(cells[penggunaan_idx]) if penggunaan_idx is not None else ""
                    perangkat_list.append({
                        "no": str(len(perangkat_list) + 1),
                        "perangkat_kerja": split_items(perangkat_raw),
                        "penggunaan_untuk_tugas": split_items(penggunaan_raw)
                    })
                break
    except Exception as e:
        print(f"❌ Gagal ekstrak perangkat kerja: {e}")
    return perangkat_list

def extract_table_after_heading(doc, heading_keywords=("tanggung jawab",), required_headers=("no.", "uraian")):
    """
    Util umum: cari heading paragraf (contains any keyword), ambil tabel pertama setelahnya
    yang headernya mengandung required_headers.
    """
    seen_heading = False
    for kind, obj in iter_block_items(doc):
        if kind == "p":
            txt = para_text(obj).lower()
            if any(k in txt for k in heading_keywords):
                seen_heading = True
        elif kind == "t" and seen_heading:
            headers = table_header_cells(obj)
            if all(h in headers for h in required_headers):
                return obj
            # kalau tabel pertama tidak cocok, lanjut cari tabel berikutnya sampai cocok
    return None

def extract_tanggung_jawab(doc):
    tanggung_list = []
    table = extract_table_after_heading(doc, ("tanggung jawab",), ("no.", "uraian"))
    if not table:
        return tanggung_list
    headers = table_header_cells(table)
    uraian_idx = headers.index("uraian")
    for r in table.rows[1:]:
        uraian_raw = extract_bullet_marked_text_cell(r.cells[uraian_idx])
        uraian_cleaned = " ".join([item.strip() for item in uraian_raw.split("|||") if item.strip()])
        tanggung_list.append({
            "no": str(len(tanggung_list) + 1),
            "uraian": uraian_cleaned
        })
    return tanggung_list

def extract_wewenang(doc):
    wewenang_list = []
    table = extract_table_after_heading(doc, ("wewenang",), ("no.", "uraian"))
    if not table:
        return wewenang_list
    headers = table_header_cells(table)
    uraian_idx = headers.index("uraian")
    for r in table.rows[1:]:
        uraian_raw = extract_bullet_marked_text_cell(r.cells[uraian_idx])
        uraian_cleaned = " ".join([item.strip() for item in uraian_raw.split("|||") if item.strip()])
        wewenang_list.append({
            "no": str(len(wewenang_list) + 1),
            "uraian": uraian_cleaned
        })
    return wewenang_list

def extract_korelasi_jabatan(doc):
    korelasi_list = []
    try:
        for table in doc.tables:
            headers = table_header_cells(table)
            if not headers:
                continue
            if ("jabatan" in headers
                and any(h in headers for h in ["unit kerja/instansi", "unit kerja", "instansi"])
                and any("dalam hal" in h for h in headers)):
                jabatan_idx = headers.index("jabatan")
                unit_idx = next((i for i, h in enumerate(headers) if "unit kerja" in h or "instansi" in h), None)
                hal_idx = next((i for i, h in enumerate(headers) if "dalam hal" in h), None)
                for r in table.rows[1:]:
                    cells = r.cells
                    jabatan_raw = extract_bullet_marked_text_cell(cells[jabatan_idx])
                    unit_raw = extract_bullet_marked_text_cell(cells[unit_idx]) if unit_idx is not None else ""
                    hal_raw = extract_bullet_marked_text_cell(cells[hal_idx]) if hal_idx is not None else ""
                    jabatan_items = " ".join([i.strip() for i in jabatan_raw.split("|||") if i.strip()])
                    unit_items = " ".join([i.strip() for i in unit_raw.split("|||") if i.strip()])
                    hal_items = split_items(hal_raw)
                    korelasi_list.append({
                        "no": str(len(korelasi_list) + 1),
                        "jabatan": jabatan_items,
                        "unit_kerja_instansi": unit_items,
                        "dalam_hal": hal_items
                    })
                break
    except Exception as e:
        print(f"❌ Gagal ekstrak korelasi jabatan: {e}")
    return korelasi_list

def extract_kondisi_lingkungan_kerja(doc):
    kondisi_list = []
    try:
        for table in doc.tables:
            headers = table_header_cells(table)
            if not headers:
                continue
            if ("aspek" in headers) and ("faktor" in headers):
                aspek_idx = headers.index("aspek")
                faktor_idx = headers.index("faktor")
                for r in table.rows[1:]:
                    aspek_raw = extract_bullet_marked_text_cell(r.cells[aspek_idx])
                    faktor_raw = extract_bullet_marked_text_cell(r.cells[faktor_idx])
                    aspek_items = " ".join([i.strip() for i in aspek_raw.split("|||") if i.strip()])
                    faktor_items = " ".join([i.strip() for i in faktor_raw.split("|||") if i.strip()])
                    kondisi_list.append({
                        "no": str(len(kondisi_list) + 1),
                        "aspek": aspek_items,
                        "faktor": faktor_items
                    })
                break
    except Exception as e:
        print(f"❌ Gagal ekstrak kondisi lingkungan kerja: {e}")
    return kondisi_list

def extract_risiko_bahaya(doc):
    risiko_list = []
    try:
        for table in doc.tables:
            headers = table_header_cells(table)
            if not headers:
                continue
            if ("nama risiko" in headers) and ("penyebab" in headers):
                risiko_idx = headers.index("nama risiko")
                penyebab_idx = headers.index("penyebab")
                for r in table.rows[1:]:
                    risiko_raw = extract_bullet_marked_text_cell(r.cells[risiko_idx])
                    penyebab_raw = extract_bullet_marked_text_cell(r.cells[penyebab_idx])
                    risiko_items = " ".join([i.strip() for i in risiko_raw.split("|||") if i.strip()])
                    penyebab_items = " ".join([i.strip() for i in penyebab_raw.split("|||") if i.strip()])
                    risiko_list.append({
                        "no": str(len(risiko_list) + 1),
                        "nama_risiko": risiko_items,
                        "penyebab": penyebab_items
                    })
                break
    except Exception as e:
        print(f"❌ Gagal ekstrak risiko bahaya: {e}")
    return risiko_list

def extract_syarat_jabatan(doc):
    result = {
        "keterampilan_kerja": [],
        "bakat_kerja": [],
        "temperamen_kerja": [],
        "minat_kerja": [],
        "upaya_fisik": [],
        "kondisi_fisik": {
            "jenis_kelamin": "",
            "umur": "",
            "tinggi_badan": "",
            "berat_badan": "",
            "postur_badan": "",
            "penampilan": "",
            "keadaan_fisik": ""
        },
        "fungsi_pekerja": []
    }

    import re

    # ---------- Helpers ----------
    def tidy(s: str) -> str:
        s = (s or "").replace("\u00A0", " ")
        s = re.sub(r"[\u0000-\u001F]", "", s)
        s = s.strip().strip(":;")
        # buang bullet/numbering di depan
        s = re.sub(r"^\s*(?:[\u2022\-\–\—\*•]+|\(?\d+[\.\)])\s*", "", s)
        return s.strip()

    def is_letter_tag(s: str) -> bool:
        """
        Tag huruf kolom-1: terima variasi 'g', 'g.', '(g)', 'g)', 'g :', 'g -', dst.
        Logika: hapus spasi & tanda baca umum, sisa tepat 1 huruf.
        """
        low = tidy(s).lower()
        core = re.sub(r"[\s\.\:\)\(\-–—]+", "", low)  # buang spasi & punct umum
        return len(core) == 1 and core.isalpha()

    def cell_items_list(cell) -> list[str]:
        items = []
        if cell is None:
            return items
        for p in cell.paragraphs:
            t = tidy(p.text)
            if not t or t in {"-", "–", "—", ":"}:
                continue
            items.append(t)
        if items:
            return items
        # fallback: split by newline
        txt = tidy(cell.text)
        return [x for x in (ln.strip() for ln in txt.splitlines()) if x]

    # prefer kolom-4, fallback kolom-3 (kalau kolom-3 bukan ":" kosong)
    def value_items_from_row(c3, c4) -> list[str]:
        v4 = cell_items_list(c4)
        if v4:
            return v4
        t3 = tidy(c3.text) if c3 is not None else ""
        if t3 and t3 != ":":
            return cell_items_list(c3)
        return []

    # khusus fungsi_pekerja: 4 -> 3 -> 2
    def value_items_for_fungsi(c2, c3, c4) -> list[str]:
        v = value_items_from_row(c3, c4)
        if v:
            return v
        return cell_items_list(c2)

    KEY_MAP = {
        "keterampilan kerja": "keterampilan_kerja",
        "bakat kerja": "bakat_kerja",
        "temperamen kerja": "temperamen_kerja",
        "minat kerja": "minat_kerja",
        "upaya fisik": "upaya_fisik",
        "kondisi fisik": "kondisi_fisik",
        "fungsi pekerja": "fungsi_pekerja",
    }
    CF_MAP = {
        "jenis kelamin": "jenis_kelamin",
        "umur": "umur",
        "tinggi badan": "tinggi_badan",
        "berat badan": "berat_badan",
        "postur badan": "postur_badan",
        "penampilan": "penampilan",
        "keadaan fisik": "keadaan_fisik",
    }

    def detect_key(label: str) -> str | None:
        low = tidy(label).lower()
        for k, v in KEY_MAP.items():
            if k in low:
                return v
        return None

    def detect_cf_field(label: str) -> str | None:
        low = tidy(label).lower()
        for k, v in CF_MAP.items():
            if k in low:
                return v
        return None

    def dedup_keep_order(seq: list[str]) -> list[str]:
        out, seen = [], set()
        for s in seq:
            t = tidy(s)
            if not t:
                continue
            if t not in seen:
                seen.add(t)
                out.append(t)
        return out

    # ---------- Pilih tabel kandidat ----------
    def table_score(tbl) -> int:
        probe = min(8, len(tbl.rows)) or 1
        if any(len(r.cells) < 4 for r in tbl.rows[:probe]):
            return -1
        score = 0
        for r in tbl.rows:
            c2 = tidy(r.cells[1].text) if len(r.cells) >= 2 else ""
            if detect_key(c2):
                score += 1
        return score

    candidate, best = None, -1
    for t in doc.tables:
        sc = table_score(t)
        if sc > best:
            best = sc
            candidate = t
    if not candidate or best <= 0:
        return result

    # ---------- Parse baris per baris ----------
    current_key: str | None = None
    in_kondisi = False

    for row in candidate.rows:
        if len(row.cells) < 4:
            continue

        c1, c2, c3, c4 = row.cells[0], row.cells[1], row.cells[2], row.cells[3]
        t1, t2 = tidy(c1.text), tidy(c2.text)

        has_letter = is_letter_tag(t1)
        key_here = detect_key(t2)
        vals_34 = value_items_from_row(c3, c4)

        # 1) Header key: ada huruf & ada label key
        if has_letter and key_here:
            current_key = key_here
            in_kondisi = (current_key == "kondisi_fisik")
            if in_kondisi:
                continue  # header kondisi_fisik umumnya tanpa nilai
            # fungsi_pekerja: 4 -> 3 -> 2
            vals = value_items_for_fungsi(c2, c3, c4) if current_key == "fungsi_pekerja" else vals_34
            if vals:
                result[current_key].extend(vals)
            continue

        # 1b) MERGED-LETTER continuation (kolom-1 masih 'g'/'g.' tapi kolom-2 kosong):
        if has_letter and (not key_here) and current_key and (current_key != "kondisi_fisik"):
            vals = value_items_for_fungsi(c2, c3, c4) if current_key == "fungsi_pekerja" else vals_34
            if vals:
                result[current_key].extend(vals)  # menangkap item pertama semisal "D0 = ..."
            continue

        # 2) Lanjutan key biasa: tanpa huruf & tanpa label key
        if (not has_letter) and (not key_here) and current_key and (current_key != "kondisi_fisik"):
            vals = value_items_for_fungsi(c2, c3, c4) if current_key == "fungsi_pekerja" else vals_34
            if vals:
                result[current_key].extend(vals)
            continue

        # 3) Mode kondisi_fisik: kolom-2 = sub-field, kolom-4/3 = nilainya (single)
        if in_kondisi and (not has_letter) and t2:
            cf_field = detect_cf_field(t2)
            if cf_field:
                val_single = " ".join(vals_34).strip()
                if val_single:
                    result["kondisi_fisik"][cf_field] = val_single
            continue

        # 4) Baris tanpa huruf tetapi ada key di kolom-2 → mulai key baru
        if (not has_letter) and key_here:
            current_key = key_here
            in_kondisi = (current_key == "kondisi_fisik")
            if in_kondisi:
                continue
            vals = value_items_for_fungsi(c2, c3, c4) if current_key == "fungsi_pekerja" else vals_34
            if vals:
                result[current_key].extend(vals)
            continue

        # 5) Penutup mode kondisi_fisik jika ada huruf tapi label kosong
        if has_letter and not key_here and in_kondisi:
            in_kondisi = False
            current_key = None
            continue

    # ---------- Fallback khusus: fungsi_pekerja masih kosong → scan ulang ----------
    if not result["fungsi_pekerja"]:
        rows = candidate.rows
        start_idx = None
        for i, r in enumerate(rows):
            if len(r.cells) < 4:
                continue
            lbl = tidy(r.cells[1].text).lower()
            if "fungsi pekerja" in lbl:
                start_idx = i
                break
        if start_idx is not None:
            for j in range(start_idx + 1, len(rows)):
                r = rows[j]
                if len(r.cells) < 4:
                    continue
                nxt_lbl = tidy(r.cells[1].text).lower()
                if any(k in nxt_lbl for k in KEY_MAP.keys() if k != "fungsi pekerja"):
                    break
                vals = value_items_for_fungsi(r.cells[1], r.cells[2], r.cells[3])  # 4 -> 3 -> 2
                if vals:
                    result["fungsi_pekerja"].extend(vals)

    # ---------- Rapikan ----------
    for k in ("keterampilan_kerja", "bakat_kerja", "temperamen_kerja",
              "minat_kerja", "upaya_fisik", "fungsi_pekerja"):
        result[k] = dedup_keep_order(result[k])

    return result


def extract_prestasi_dan_kelas(doc):
    lines = [para_text(p) for p in doc.paragraphs if para_text(p)]
    prestasi = "---"
    kelas = "---"
    for i, line in enumerate(lines):
        lower = line.lower().strip()
        if "prestasi yang diharapkan" in lower and i + 1 < len(lines):
            prestasi = lines[i + 1].strip()
        elif "kelas jabatan" in lower and i + 1 < len(lines):
            kelas = lines[i + 1].strip()
    return prestasi, kelas

# -------------------- ORKESTRATOR --------------------

def extract_info(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".docx":
        doc, lines = read_docx(file_path)
    elif ext == ".doc":
        doc, lines = read_doc(file_path)  # konversi via LibreOffice, lalu python-docx
    else:
        raise ValueError("File tidak didukung: " + file_path)

    prestasi, kelas = extract_prestasi_dan_kelas(doc)
    return {
        "file": os.path.basename(file_path),
        "nama_jabatan": extract_line_value("NAMA JABATAN", lines),
        "kode_jabatan": extract_line_value("KODE JABATAN", lines),
        "unit_kerja": extract_unit_kerja(lines),
        "ikhtisar_jabatan": extract_block("IKHTISAR JABATAN", "KUALIFIKASI JABATAN", lines),
        "kualifikasi_jabatan": extract_kualifikasi(doc),
        "tugas_pokok": extract_tugas_pokok(doc),
        "hasil_kerja": extract_hasil_kerja(doc),
        "bahan_kerja": extract_bahan_kerja(doc),
        "perangkat_kerja": extract_perangkat_kerja(doc),
        "tanggung_jawab": extract_tanggung_jawab(doc),
        "wewenang": extract_wewenang(doc),
        "korelasi_jabatan": extract_korelasi_jabatan(doc),
        "kondisi_lingkungan_kerja": extract_kondisi_lingkungan_kerja(doc),
        "risiko_bahaya": extract_risiko_bahaya(doc),
        "syarat_jabatan": extract_syarat_jabatan(doc),
        "prestasi_yang_diharapkan": prestasi,
        "kelas_jabatan": kelas
    }

# -------------------- CLI --------------------

if __name__ == "__main__":
    import sys
    file_path = globals().get("__file_path__", None)
    if not file_path and len(sys.argv) >= 2:
        file_path = sys.argv[1]

    if file_path:
        try:
            data = extract_info(file_path)
            print(json.dumps(data, ensure_ascii=False))
        except Exception as e:
            print(f"❌ Error: {str(e)}", file=sys.stderr)
            sys.exit(1)
    else:
        print("❌ Argumen tidak lengkap: butuh file .doc/.docx", file=sys.stderr)
        sys.exit(1)
